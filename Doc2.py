#!/bin/python
import docx
from docx.shared import Pt
#from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

source_doc = docx.Document('InitialDoc.docx')
new_doc = docx.Document()

dash = '–' 


font_style_heading = new_doc.styles.add_style('HeadingStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
font_style_heading.font.name = 'Arial'
font_style_heading.font.size = Pt(16)

font_style_normal = new_doc.styles.add_style('NormalStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
font_style_normal.font.name = 'Comic Sans MS'
font_style_normal.font.size = Pt(12)

font_style_empty = new_doc.styles.add_style('EmptyStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
font_style_empty.font.name = 'Calibri'
font_style_empty.font.size = Pt(10)


for paragraph in source_doc.paragraphs:
    string = str(paragraph.text)
    if string.count(dash)!= 0 and string != '\r': #было /n, но не работало, сейчас стоит /r для переноси строки
        new_paragraph = new_doc.add_paragraph(paragraph.text, style='NormalStyle')

    elif string == "\r": #было /n, но не работало, сейчас стоит /r для переноси строки (x2)
        new_paragraph = new_doc.add_paragraph(paragraph.text, style='EmptyStyle')

    elif string != "\r": #было /n, но не работало, сейчас стоит /r для переноси строки (x3)
        new_paragraph = new_doc.add_paragraph(paragraph.text, style='HeadingStyle')
    


new_doc.save('./Result/NewDoc.docx') #NewDoc.docx  ./Result/
